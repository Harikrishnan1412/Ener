import pandas as pd
import os
import docx
from docx import Cm, Pt, Inches
#from docx.shared import Cm, Pt, Inches
from docx import WD_ALIGN_PARAGRAPH
#from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter 
from tkinter import *
from tkinter import ttk
import tkinter.font as tkfont
#from tkcalendar import Calendar, DateEntry
from datetime import datetime
from tkinter.filedialog import askopenfilename
from PIL import ImageTk, Image
from tkinter import messagebox
from datetime import timedelta, date
from datetime import datetime
from tkcalendar import Calendar, DateEntry
#import babel.numbers
import numpy
#import xlwt
#import xlrd
#from xlutils.copy import copy
from collections import OrderedDict

def show():
    messagebox.showinfo("Success","Document Created")
def show2():
    messagebox.showinfo("Success","Hall Chart and Attendance Cover Created")
def alert():
    messagebox.showerror("Error","Slot and Course Code Not Matching")

#info page
def info1():
    root_a = Toplevel()
    #root_a = tkinter.Tk()
    root_a.configure(background='lavender')
    root_a.title("Info")
    root_a.geometry("650x550")
    #root_a.iconbitmap('tce.ico')
    
    # Heading
    w2 = tkinter.Label(root_a,  text="TEAM", fg="black", bg="lavender")
    w2.config(font=("Elephant", 20))
    w2.place(x=275,y=5)


    #Images

    load= (Image.open("Photos\sabari.png"))
    load = load.resize((90, 90), Image.ANTIALIAS)
    render = ImageTk.PhotoImage(load)
    img = Label(root_a, image=render)
    img.place(x=20, y=50)

    load1= (Image.open("Photos\hari.png"))
    load1 = load1.resize((90, 90), Image.ANTIALIAS)
    render1 = ImageTk.PhotoImage(load1)
    img1 = Label(root_a, image=render1)
    img1.place(x=320, y=50)

    load2= (Image.open("Photos\Vaibhaav.png"))
    load2 = load2.resize((90, 90), Image.ANTIALIAS)
    render2 = ImageTk.PhotoImage(load2)
    img2 = Label(root_a, image=render2)
    img2.place(x=20, y=150)

    load3= (Image.open("Photos\kabilesh.png"))
    load3 = load3.resize((90, 90), Image.ANTIALIAS)
    render3 = ImageTk.PhotoImage(load3)
    img3 = Label(root_a, image=render3)
    img3.place(x=320, y=150)



    # labels
    NameLb = tkinter.Label(root_a, text="B.Sabari Shri Varshan", fg="black", bg="lavender")
    NameLb.config(font=("Aharoni", 10))
    NameLb.place(x=120,y=70)

    MailLb = tkinter.Label(root_a, text="sabarishrivarshan@gmail.com", fg="black", bg="lavender")
    MailLb.config(font=("Aharoni", 10))
    MailLb.place(x=120,y=100)

    NameLb1 = tkinter.Label(root_a, text="G.Harikrishnan", fg="black", bg="lavender")
    NameLb1.config(font=("Aharoni", 10))
    NameLb1.place(x=420,y=70)

    MailLb1 = tkinter.Label(root_a, text="harikrishnan14122000@gmail.com", fg="black", bg="lavender")
    MailLb1.config(font=("Aharoni", 10))
    MailLb1.place(x=420,y=100)

    NameLb2 = tkinter.Label(root_a, text="T.L.Vaibhaav Ram", fg="black", bg="lavender")
    NameLb2.config(font=("Aharoni", 10))
    NameLb2.place(x=120,y=170)

    MailLb2 = tkinter.Label(root_a, text="vaibhaavram01@gmail.com", fg="black", bg="lavender")
    MailLb2.config(font=("Aharoni", 10))
    MailLb2.place(x=120,y=200)

    NameLb3 = tkinter.Label(root_a, text="A.Kabilesh", fg="black", bg="lavender")
    NameLb3.config(font=("Aharoni", 10))
    NameLb3.place(x=420,y=170)

    MailLb3 = tkinter.Label(root_a, text="Kabileshashok@gmail.com", fg="black", bg="lavender")
    MailLb3.config(font=("Aharoni", 10))
    MailLb3.place(x=420,y=200)

    FootLb = tkinter.Label(root_a, text="Version 1.0.4, Designed for EEE TCE", fg="black", bg="lavender")
    FootLb.config(font=("Aharoni", 8))
    FootLb.place(relx = 0.0,rely = 1.0,anchor ='sw')

    DesLb = tkinter.Label(root_a, text="Packages used and their Version", fg="red", bg="lavender")
    DesLb.config(font=("Aharoni", 10))
    DesLb.place(x=20,y=470)

    DesLb1 = tkinter.Label(root_a, text="Python 3.7.6, tkinter version 8.6, python-docx 0.8.11, pandas 1.0.5", fg="blue", bg="lavender")
    DesLb1.config(font=("Aharoni", 10))
    DesLb1.place(x=20,y=490)

    #Button
    root_a.ext = tkinter.Button(root_a, text="   Back   ", command=root_a.destroy,bg="bisque",fg="blue")
    root_a.ext.place(x=550,y=520)
    root_a.ext.config(font=("Aharoni", 12))

    root_a.mainloop()


#attendence page
def att():
    subcode = []
    def subject_code():
            j = -1
            a= root.subcd.get()
            subcode.append(a)
            code = subcode[j]
            roll = []
            file = askopenfilename()
            subject_code = list(pd.read_excel(file, sheet_name=code,usecols="B").columns.values)
            regno1 = pd.read_excel(file, sheet_name=code,usecols="A").values
            regno2 = regno1.flatten()
            name1 = pd.read_excel(file, sheet_name=code,usecols="B").values
            name2 = name1.flatten()
            res = dict(zip(regno2[3:], name2[3:]))
            res1 = OrderedDict(sorted(res.items()))
            #res1 = dict(sorted(res.items(), key=lambda item: item[1]))

            # Import docx NOT python-docx


            # Create an instance of a word document
            doc = docx.Document('Default.docx')

            # Add a Title to the document

            doc_para = doc.add_paragraph(subject_code[0])
            doc_para = doc.add_paragraph('Staff Info: '+ name2[0])

            # Table data in a form of list

            # Creating a table object
            table = doc.add_table(rows=1, cols=0)
            doc_para = doc.add_paragraph('          ')
            table1 = doc.add_table(rows=4, cols=0)

            first_column_width = 3
            second_column_width = 5
            third_column_width = 13
            fourth_column_width = 3
            table.add_column(Cm(first_column_width))
            table.add_column(Cm(second_column_width))
            table.add_column(Cm(third_column_width))
            table.add_column(Cm(fourth_column_width))
            table.add_column(Cm(fourth_column_width))
            table.add_column(Cm(fourth_column_width))
            table1.add_column(Cm(fourth_column_width))
            table1.add_column(Cm(fourth_column_width))
            table1.add_column(Cm(fourth_column_width))
            table1.add_column(Cm(fourth_column_width))
            

            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row1 = table1.rows[0].cells
            col = table1.columns[0].cells
            row[0].text = 'S.No'
            row[1].text = 'RegNo'
            row[2].text = 'Name'
            row[3].text = 'CAT 1'
            row[4].text = 'CAT 2'
            row[5].text = 'CAT 3'
            row1[0].text = 'CAT'
            row1[1].text = 'Present'
            row1[2].text = 'Absent'
            row1[3].text = 'Total'
            col[1].text = 'I'
            col[2].text = 'II'
            col[3].text = 'III'
            i = 1
            # Adding data from the list to the table
            for j in res1:

                    # Adding a row and then adding data in it.
                    row = table.add_row().cells
                    row[0].text = str(i)
                    row[0].paragraphs[0].runs[0].font.size = Pt(10)
                    row[1].text = j
                    row[1].paragraphs[0].runs[0].font.size = Pt(10)
                    row[2].text = res[j]
                    row[2].paragraphs[0].runs[0].font.size = Pt(10)
                    i = i+1
            # Adding style to a table
            table.style = 'Table Grid'
            table1.style = 'Table Grid'

            # Now save the document to a location
            new_path = 'Attendance Document\\'
            name = code+' Attendance.docx'
            if(not os.path.exists(new_path)):
                    os.makedirs(new_path)
            doc.save(new_path+name)
            show()

    root = tkinter.Tk()
    root.configure(background='lavender')
    root.title("Subject Code")
    root.geometry("660x250")
    #root.iconbitmap('tce.ico')

    # Heading
    w2 = tkinter.Label(root,  text="ATTENDANCE DOCUMENT", fg="black", bg="lavender")
    w2.config(font=("Elephant", 30))
    w2.grid(row=1, column=0, columnspan=2, padx=10)

    # labels
    NameLb = tkinter.Label(root, text="Enter the Subject Code", fg="black", bg="lavender")
    NameLb.config(font=("Aharoni", 15))
    NameLb.grid(row=6, column=0, pady=15)

    FootLb = tkinter.Label(root, text="Version 1.0.4, Designed for EEE TCE", fg="black", bg="lavender")
    FootLb.config(font=("Aharoni", 8))
    FootLb.place(relx = 0.0,rely = 1.0,anchor ='sw')

    #entries
    root.subcd = tkinter.Entry(root, textvariable=subcode)
    root.subcd.grid(row=6, column=1)

    #button
    root.star = tkinter.Button(root, text="   Proceed   ", command=subject_code,bg="bisque",fg="blue")
    root.star.grid(row=8, column=0,padx=0)
    root.star.config(font=("Aharoni", 12))

    root.star = tkinter.Button(root, text="   Exit   ", command=root.destroy,bg="bisque",fg="blue")
    root.star.grid(row=8, column=1,padx=0)
    root.star.config(font=("Aharoni", 12))
    root.mainloop()

#Hall chart
def hallchart1():
    #Define Function to get values from tkinter
    hall=[]
    hall1 = []
    sub1=[]
    sub2=[]
    date_s = []
    date_e = []
    hall_no = []
    hall_no1 = []
    U_P = []
    ns = []
    sem = []
    slot = [0]
    core_course = [0]
    sem1 = []
    slot1 = [0]
    core_course1 = [0]
    t = []


#elective tkinter page function start
    def ele_co():
        def show():
                messagebox.showinfo("Success","Document Created")
        def showele():
            messagebox.showinfo("Info","Select the elective List")
        def showele1():
            messagebox.showinfo("Info","Select the Student List")
            
        coresc = []
        ehall1 = []
        ehall2 = []
        et = []
        cat=[]
        x= []
        def elective():
            c11 = root_e.coresc.get()
            coresc.append(c11)
            core_cd = coresc[-1]
            c12 = root_e.ehall1.get()
            ehall1.append(c12)
            ehall11 = ehall1[-1]
            c13 = root_e.ehall2.get()
            ehall2.append(c13)
            ehall12 = ehall2[-1]
            c14 = root_e.et.get()
            et.append(c14)
            et1 = et[-1]
            c15 = root_e.cat.get()
            cat.append(c15)
            cat1 = cat[-1]
            print(ehall11)
            print(ehall12)
            showele()
            filee = askopenfilename()
            subject_codes1 = pd.read_excel(filee,usecols="A").values
            subject_codes = subject_codes1.flatten()
            showele1()
            filees = askopenfilename()
            print("Files are collected")
            for sc in subject_codes:
                regnoe1 = pd.read_excel(filees, sheet_name=sc,usecols="A").values
                regnoe3 = regnoe1.flatten()
                regnoe2 = regnoe3[3:]
                regnoe2.sort()
                x1 = []
                for i in regnoe2:
                    if('E' in i):
                        x1.append(i)
                x1.append(sc)
                x.append(x1)
            print("Lsit ready")
           # print("Select the core subject student list")
           # fileecs = askopenfilename()
            #print("Enter the core subject code")
            #core_cd = input()
            regnoce1 = pd.read_excel(filees, sheet_name=core_cd,usecols="A").values
            regnoce3 = regnoce1.flatten()
            regnoce2 = regnoce3[3:]
            name1 = pd.read_excel(filees, sheet_name=core_cd,usecols="B").values
            name = name1.flatten()
            dictm1 = dict(zip(regnoce2,name[3:]))
            dictm = dict(sorted(dictm1.items(), key=lambda item: item[0]))
            print("PRint dictm")
            print(dictm)
            regnoce2 = list(dictm.keys())
            print("     ----     ")
            print(regnoce2)
            print(len(regnoce2))
            print("    ------    ")
            dict1 = {}
            dictn = {}
            dictn1 = {}
            for b in x:
                c = 0
                x2 = []
                for c1 in b[:-1]:
                    if(c1 in regnoce2[:35]):
                        c = c+1
                        #print(c1)
                        x2.append(c1)
                dictn[b[-1]] = x2
                dict1[b[-1]] = c
            print(dict1)
            dict2 = {}
            for b in x:
                c = 0
                x3 = []
                for c1 in b[:-1]:
                    if(c1 in regnoce2[35:]):
                        c = c+1
                        #print(c1)
                        x3.append(c1)
                dictn1[b[-1]] = x3
                dict2[b[-1]] = c
            print(dict2)
            print(dictn)
            print(dictn1)
            doc = docx.Document('Default with Footer.docx')
            doc_para6 = doc.add_paragraph('                                                                      '+et1+' Elective Student Count                                                             ')
            doc_para3 = doc.add_paragraph('Hall number : '+ehall11+'                ')
            #Total count for elective count in First Hall
            table = doc.add_table(rows=1,cols=2)
            doc_para = doc.add_paragraph('          ')
            doc_para1 = doc.add_paragraph('          ')
            doc_para4 = doc.add_paragraph('Hall number : '+ehall12+'                ')
            #total elective count for second Hall
            table1 = doc.add_table(rows=1,cols=2)
            doc_para3 = doc.add_paragraph('          ')
            #first Hall count table
            row = table.rows[0].cells
            row[0].text = 'Course code'
            row[1].text = 'Number of Studeents'
            #second Hall count table
            row1 = table1.rows[0].cells
            row1[0].text = 'Course code'
            row1[1].text = 'Number of Studeents'
            #forloop for count tables in Hall 1
            for i in dict1:
                row = table.add_row().cells
                row[0].text = i
                row[0].paragraphs[0].runs[0].font.size = Pt(10)
                row[1].text = str(dict1[i])
                row[1].paragraphs[0].runs[0].font.size = Pt(10)
            for j in dictn:
                doc_para5 = doc.add_paragraph('Elective course code : '+j)
                table2 = doc.add_table(rows=1,cols=3)
                table2.style = 'Table Grid'
                row2 = table2.rows[0].cells
                row2[0].text = 'Reg No'
                row2[1].text = 'Name'
                row2[2].text = 'Hall No'
                for k in dictn[j]:
                    row3 = table2.add_row().cells
                    row3[0].text = k
                    row3[1].text = dictm[k]
                    row3[2].text = ehall11
                doc_para1 = doc.add_paragraph('          ')
            for a in dict2:
                row1 = table1.add_row().cells
                row1[0].text = a
                row1[0].paragraphs[0].runs[0].font.size = Pt(10)
                row1[1].text = str(dict2[a])
                row1[1].paragraphs[0].runs[0].font.size = Pt(10)
            for b in dictn1:
                doc_para5 = doc.add_paragraph('Elective course code : '+b)
                table3 = doc.add_table(rows=1,cols=3)
                table3.style = 'Table Grid'
                row4 = table3.rows[0].cells
                row4[0].text = 'Reg No'
                row4[1].text = 'Name'
                row4[2].text = 'Hall No'
                for c in dictn1[b]:
                    row5 = table3.add_row().cells
                    row5[0].text = c
                    row5[1].text = dictm[c]
                    row5[2].text = ehall12
                doc_para1 = doc.add_paragraph('          ')
            table.style = 'Table Grid'
            table1.style = 'Table Grid'

            # Now save the document to a location
            new_path = 'Elective Document\\'
            name = cat1+' '+et1+' '+ehall11+','+ehall12+' Hall elective numbers.docx'
            if(not os.path.exists(new_path)):
                os.makedirs(new_path)
            doc.save(new_path+name)
            show()
            print("Completed")

            
        root_e = Toplevel()
        root_e.geometry("500x350")

        # Heading
        rw2 = tkinter.Label(root_e,  text="ELECTIVE STUDENT COUNT", fg="black", bg="lavender")
        rw2.config(font=("Elephant", 20))
        rw2.place(x=25,y=5)

        #Labels
        Lb = tkinter.Label(root_e, text="Enter Core subject code :", fg="black", bg="lavender")
        Lb.config(font=("Aharoni", 10))
        Lb.place(x=20,y=70)

        Lb1 = tkinter.Label(root_e, text="Enter Hall number 1 :", fg="black", bg="lavender")
        Lb1.config(font=("Aharoni", 10))
        Lb1.place(x=20,y=100)

        Lb2 = tkinter.Label(root_e, text="Enter Hall number 2 :", fg="black", bg="lavender")
        Lb2.config(font=("Aharoni", 10))
        Lb2.place(x=20,y=130)

        Lb3 = tkinter.Label(root_e, text="Enter elective type :", fg="black", bg="lavender")
        Lb3.config(font=("Aharoni", 10))
        Lb3.place(x=20,y=160)

        Lb3 = tkinter.Label(root_e, text="Enter Cat details :", fg="black", bg="lavender")
        Lb3.config(font=("Aharoni", 10))
        Lb3.place(x=20,y=190)

        #button
        root_e.submit = tkinter.Button(root_e,text = "  Proceed  ",command=elective,bg="bisque",fg="blue")
        root_e.submit.place(x=250,y=240)
        root_e.submit.config(font=("Aharoni", 12))

        #entries
        root_e.coresc = tkinter.Entry(root_e, textvariable=coresc)
        root_e.coresc.place(x=250,y=70)

        root_e.ehall1 = tkinter.Entry(root_e, textvariable=ehall1)
        root_e.ehall1.place(x=250,y=100)

        root_e.ehall2 = tkinter.Entry(root_e, textvariable=ehall2)
        root_e.ehall2.place(x=250,y=130)

        root_e.et = tkinter.Entry(root_e, textvariable=et)
        root_e.et.place(x=250,y=160)

        root_e.cat = tkinter.Entry(root_e, textvariable=cat)
        root_e.cat.place(x=250,y=190)

        root_e.mainloop()




#elective tkinter page function end
    
    def proceed():
        #title append
        print(n.get())
        t.append(n.get())#title add array
        print(t)
        #date append
        t11 = cal.get_date()
        t12 = cal1.get_date()
        format = "%d.%m.%y"
        t11=t11.strftime(format)
        t12=t12.strftime(format)
        date_s.append(t11) #start date input array
        date_e.append(t12)#end date input array
        print(date_s,date_e)
        #Hall append
        a=root1.hallEn.get()
        hall_no.append(a) #first Hall number
        print(hall_no)
        a11 = root1.hallEn1.get()
        hall_no1.append(a11) #second Hall number
        print(hall_no1)
        #UG or PG
        U_P.append(n1.get())#UG/PG enter array
        print(U_P)
        #number of sem
        ns.append(n2.get())#Number of semester chossen Enter array
        print(ns)

        #based on number of sem choosen below if case work

        if(ns[-1]=='1'):  #choices if only one semester is choosen
            print("1")
            #1 sem detail
            sem.append(n3.get())#one sem choosen sem detail array
            print(sem) 
            #1 sem slot
            slot.append(n4.get())#one sem shoosen slot detail array
            print(slot)
            #1 sem core course code
            a1=root1.sub1En.get() #one sem choosen subject code choose array
            core_course.append(a1)
            print(core_course)
        elif(ns[-1]=='2'): #choices if two semester is choosen
            print("2")
            #1 sem detail
            sem.append(n3.get())
            print(sem)
            #1 sem slot
            slot.append(n4.get())
            print(slot)
            #1 sem core course code
            a1=root1.sub1En.get()
            core_course.append(a1)
            print(core_course)
            #2 sem deatil
            sem1.append(n5.get())
            print(sem1)
            #2 sem slot
            slot1.append(n6.get())
            print(slot1)
            #2 sem core course code
            a2=root1.sub2En.get()
            core_course1.append(a2)

        if(((slot[-1]=='A' and core_course[-1][-1]!='1')or(slot[-1]=='B' and core_course[-1][-1]!='2'))or((slot1[-1]=='A' and core_course1[-1][-1]!='1')or(slot1[-1]=='B' and core_course1[-1][-1]!='2'))):
            print(slot[-1])
            print(core_course[-1][-1])
            alert()
            return

        #excel reading
        if(ns[-1] =='1'):
            file = askopenfilename() #excel file saving variable name
            regno = pd.read_excel(file, sheet_name=core_course[-1],usecols="A").values
            regno1 = regno.flatten()
            name = pd.read_excel(file, sheet_name=core_course[-1],usecols="B").values
            name1 = name.flatten()
            dic = dict(zip(regno1[3:], name1[3:]))
            dic1 = dict(sorted(dic.items(), key=lambda item: item[0]))
            room1 = dict(list(dic1.items())[:35])
            room2 = dict(list(dic1.items())[35:])
            print("ROOM 1")
            print(room1)
            print("ROOM 2")
            print(room2)
        elif(ns[-1]=='2'):
            file = askopenfilename()
            regno = pd.read_excel(file, sheet_name=core_course[-1],usecols="A").values
            regno1 = regno.flatten()
            name = pd.read_excel(file, sheet_name=core_course[-1],usecols="B").values
            name1 = name.flatten()
            dic = dict(zip(regno1[3:], name1[3:]))
            dic1 = dict(sorted(dic.items(), key=lambda item: item[0]))
            room1 = dict(list(dic1.items())[:35])
            room2 = dict(list(dic1.items())[35:])
            """print("ROOM 1")
            print(room1)
            print("ROOM 2")
            print(room2)"""
            regno2 = pd.read_excel(file, sheet_name=core_course1[-1],usecols="A").values
            regno3 = regno2.flatten()
            name2 = pd.read_excel(file, sheet_name=core_course1[-1],usecols="B").values
            name3 = name2.flatten()
            dic2 = dict(zip(regno3[3:], name3[3:]))
            dic3 = dict(sorted(dic2.items(), key=lambda item: item[0]))
            room3 = dict(list(dic3.items())[:35])
            room4 = dict(list(dic3.items())[35:])
            

        #document creation using default
        doc = docx.Document('Default with Footer.docx')
        doc_para1 = doc.add_paragraph(t[-1]+' - Hall Chart')
        doc_para1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        #adding tables
        table = doc.add_table(rows=3, cols=2)
        doc_para = doc.add_paragraph('          ')
        if(ns[-1]=='1'):
            doc_para = doc.add_paragraph(hall_no[-1]+" ("+str(len(room1))+") - Sem "+sem[-1]+" ("+str(len(room1))+")  "+slot[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1]=='2'):
            doc_para = doc.add_paragraph(hall_no[-1]+" ("+str(len(room1)+len(room3))+") - Sem "+sem[-1]+" ("+str(len(room1))+")  "+slot[-1]+" - Sem "+sem1[-1]+" ("+str(len(room3))+") "+slot1[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        table1 = doc.add_table(rows=10, cols=10)
        table1.allow_autofit = False

        table1.cell(0,0).merge(table1.cell(0,1))
        table1.cell(0,2).merge(table1.cell(0,3))
        table1.cell(0,4).merge(table1.cell(0,5))
        table1.cell(0,6).merge(table1.cell(0,7))
        table1.cell(0,8).merge(table1.cell(0,9))

        #adding heading for first table
        row = table.rows[0].cells
        row[0].add_paragraph('Degree & Department: B.E - EEE')
        row[1].add_paragraph('Slot - I: 09.30 A.M to 11.00 A.M')
        row = table.rows[1].cells
        row[0].add_paragraph('Date: '+date_s[-1]+' - '+date_e[-1])
        row[1].add_paragraph('Slot - II: 11.30 A.M to 01.00 P.M')
        row = table.rows[2].cells
        row[1].add_paragraph('Slot - III: 02.30 P.M to 4.00 P.M')


        #adding heading
        row = table1.rows[0].cells
        p=row[0].add_paragraph('Column 1')
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[2].add_paragraph('Column 2')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[4].add_paragraph('Column 3')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[6].add_paragraph('Column 4')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[8].add_paragraph('Column 5')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER

        row = table1.rows[1].cells
        if(ns[-1] == '1'):
            p=row[0].add_paragraph('Sem '+sem[-1])
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem[-1])
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem[-1])
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem[-1])
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem[-1])
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1] == '2'):
            for i in range(10):
                if(i%2==0):
                    p=row[i].add_paragraph('Sem '+sem[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p=row[i].add_paragraph('Sem '+sem1[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

                    
        """    p=row[0].add_paragraph('Sem '+sem)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem)
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem)
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem)
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem)
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p=row[1].add_paragraph('Sem '+sem1)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[3].add_paragraph('Sem '+sem1)
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[5].add_paragraph('Sem '+sem1)
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[7].add_paragraph('Sem '+sem1)
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[9].add_paragraph('Sem '+sem1)
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER"""

        #limit of room check
        if(ns[-1] == '1'):
            flag = 0
            rn = list(room1.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(0,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn[k]
                    col[j].width=Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn)):
                        flag =1
                        break
        elif(ns[-1] == '2'):
            flag = 0
            rn = list(room1.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(0,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn[k]
                    col[j].width = Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn)):
                        flag =1
                        break
            flag = 0
            rn2 = list(room3.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(1,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn2[k]
                    col[j].width = Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn2)):
                        flag =1
                        break
            
             

        table.style = 'Table Grid'
        table1.style = 'Table Grid'

        #go for new page
        doc.add_page_break()
        doc_para1 = doc.add_paragraph(t[-1]+' - Hall Chart')
        doc_para1.alignment=WD_ALIGN_PARAGRAPH.CENTER

        table2 = doc.add_table(rows=3, cols=2)

        #adding heading for first table
        row = table2.rows[0].cells
        row[0].add_paragraph('Degree & Department: B.E - EEE')
        row[1].add_paragraph('Slot - I: 09.30 A.M to 11.00 A.M')
        row = table2.rows[1].cells
        row[0].add_paragraph('Date: '+date_s[-1]+' - '+date_e[-1])
        row[1].add_paragraph('Slot - II: 11.30 A.M to 01.00 P.M')
        row = table2.rows[2].cells
        row[1].add_paragraph('Slot - III: 02.30 P.M to 4.00 P.M')

        doc_para = doc.add_paragraph('          ')
        if(ns[-1]=='1'):
            doc_para = doc.add_paragraph(hall_no1[-1]+"  ("+str(len(room2))+") - Sem "+sem[-1]+" ("+str(len(room2))+") - "+slot[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1]=='2'):
            doc_para = doc.add_paragraph(hall_no1[-1]+"  ("+str(len(room2)+len(room4))+") - Sem "+sem[-1]+" ("+str(len(room2))+") "+slot[-1]+" - Sem "+sem1[-1]+" ("+str(len(room4))+")  "+slot1[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER

        table3 = doc.add_table(rows=10, cols=10)
        table3.allow_autofit = False

        table3.cell(0,0).merge(table3.cell(0,1))
        table3.cell(0,2).merge(table3.cell(0,3))
        table3.cell(0,4).merge(table3.cell(0,5))
        table3.cell(0,6).merge(table3.cell(0,7))
        table3.cell(0,8).merge(table3.cell(0,9))

        #adding heading
        row = table3.rows[0].cells
        p=row[0].add_paragraph('Column 1')
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[2].add_paragraph('Column 2')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[4].add_paragraph('Column 3')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[6].add_paragraph('Column 4')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[8].add_paragraph('Column 5')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER

        if(ns[-1] == '1'):
            row = table3.rows[1].cells
            p=row[0].add_paragraph('Sem '+sem[-1])
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem[-1])
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem[-1])
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem[-1])
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem[-1])
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1] == '2'):
            row = table3.rows[1].cells
            for i in range(10):
                if(i%2==0):
                    p=row[i].add_paragraph('Sem '+sem[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p=row[i].add_paragraph('Sem '+sem1[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

        #table 3 filling
        if(ns[-1] == '1'):
            flag1 = 0
            rn1 = list(room2.keys())
            if(len(room2)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(0,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn1[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn1)):
                        flag1 =1
                        break
        elif(ns[-1] == '2'):
            flag1 = 0
            rn1 = list(room2.keys())
            if(len(room2)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(0,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn1[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn1)):
                        flag1 =1
                        break
            flag1 = 0
            rn3 = list(room4.keys())
            if(len(room4)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(1,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn3[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn3)):
                        flag1 =1
                        break
            
        table2.style = 'Table Grid'
        table3.style = 'Table Grid'

        new_path = 'Hallchart Document\\'
        name = str(hall_no[-1])+" "+str(hall_no1[-1])+" "+str(t[-1])+" ("+str(date_s[-1])+") "+"Hall chart.docx"
        if(not os.path.exists(new_path)):
                os.makedirs(new_path)
        doc.save(new_path+name)


        #Attendence Cover code join here
        
        #document creation using default
        doc = docx.Document('Default with Footer.docx')
        doc_para1 = doc.add_paragraph(t[-1]+' - Hall Chart and Attendance Cover')
        doc_para1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        #adding tables
       # table = doc.add_table(rows=3, cols=2)
        #doc_para = doc.add_paragraph('          ')
        if(ns[-1]=='1'):
            doc_para = doc.add_paragraph(hall_no[-1]+" ("+str(len(room1))+") - Sem "+sem[-1]+" ("+str(len(room1))+")  "+slot[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1]=='2'):
            doc_para = doc.add_paragraph(hall_no[-1]+" ("+str(len(room1)+len(room3))+") - Sem "+sem[-1]+" ("+str(len(room1))+")  "+slot[-1]+" - Sem "+sem1[-1]+" ("+str(len(room3))+") "+slot1[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        table1 = doc.add_table(rows=10, cols=10)
        table1.allow_autofit = False

        table1.cell(0,0).merge(table1.cell(0,1))
        table1.cell(0,2).merge(table1.cell(0,3))
        table1.cell(0,4).merge(table1.cell(0,5))
        table1.cell(0,6).merge(table1.cell(0,7))
        table1.cell(0,8).merge(table1.cell(0,9))

        """#adding heading for first table
        row = table.rows[0].cells
        row[0].add_paragraph('Degree & Department: B.E - EEE')
        row[1].add_paragraph('Slot - I: 09.30 A.M to 11.00 A.M')
        row = table.rows[1].cells
        row[0].add_paragraph('Date: '+date_s[-1]+' - '+date_e[-1])
        row[1].add_paragraph('Slot - II: 11.30 A.M to 01.00 P.M')
        row = table.rows[2].cells
        row[1].add_paragraph('Slot - III: 02.30 P.M to 4.00 P.M')"""


        #adding heading
        row = table1.rows[0].cells
        p=row[0].add_paragraph('Column 1')
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[2].add_paragraph('Column 2')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[4].add_paragraph('Column 3')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[6].add_paragraph('Column 4')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[8].add_paragraph('Column 5')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER

        row = table1.rows[1].cells
        if(ns[-1] == '1'):
            p=row[0].add_paragraph('Sem '+sem[-1])
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem[-1])
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem[-1])
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem[-1])
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem[-1])
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1] == '2'):
            for i in range(10):
                if(i%2==0):
                    p=row[i].add_paragraph('Sem '+sem[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p=row[i].add_paragraph('Sem '+sem1[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

                    
        """    p=row[0].add_paragraph('Sem '+sem)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem)
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem)
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem)
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem)
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p=row[1].add_paragraph('Sem '+sem1)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[3].add_paragraph('Sem '+sem1)
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[5].add_paragraph('Sem '+sem1)
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[7].add_paragraph('Sem '+sem1)
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[9].add_paragraph('Sem '+sem1)
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER"""

        #limit of room check
        if(ns[-1] == '1'):
            flag = 0
            rn = list(room1.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(0,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn[k]
                    col[j].width=Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn)):
                        flag =1
                        break
        elif(ns[-1] == '2'):
            flag = 0
            rn = list(room1.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(0,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn[k]
                    col[j].width = Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn)):
                        flag =1
                        break
            flag = 0
            rn2 = list(room3.keys())
            """if(len(room1)==40):
                lim = 10
            else:
                lim = 9"""
            k = 0
            for i in range(1,10,2):
                col = table1.columns[i].cells
                if(flag):
                    break
                for j in range(2,9):
                    col[j].text = rn2[k]
                    col[j].width = Cm(1.0)
                    col[j].paragraphs[0].runs[0].font.size = Pt(10)
                    k = k+1
                    if(k>=len(rn2)):
                        flag =1
                        break
            
             

        #table.style = 'Table Grid'
        table1.style = 'Table Grid'

        def daterange(date1, date2):
            for n in range(int((date2 - date1).days)+1):
                yield date1 + timedelta(n)
        date1 = datetime.strptime(date_s[-1], '%d.%m.%y')
        date2 = datetime.strptime(date_e[-1], '%d.%m.%y')

        weekdays = [5, 6]
        dates = []
        for dt in daterange(date1, date2):
            if dt.weekday() not in weekdays:                    # to print only the weekdates
                dates.append(dt.strftime("%d.%m.%Y"))
        r1 = len(dates)*4
        doc_para2 = doc.add_paragraph("           ")
        #attendence cover
        
        table11 = doc.add_table(rows=r1+1, cols=8)

        #Adding height
        for row in table11.rows:
            row.height = Cm(0.1)
        for cell in table11.columns[0].cells:
            cell.width = Cm(0.8)
        #row heading
        row = table11.rows[0].cells
        p=row[0].add_paragraph('Date')
        #p.paragraphs[0].runs[0].font.size = Pt(10)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[1].add_paragraph('Session')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[2].add_paragraph('Semester')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[3].add_paragraph('Reg.No.of.Absentees')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[4].add_paragraph('Absent')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p5=row[5].add_paragraph('Present')
        p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p6=row[6].add_paragraph('Total')
        p6.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p7=row[7].add_paragraph('Signature')
        p7.alignment=WD_ALIGN_PARAGRAPH.CENTER
        #table.cell(1,2).merge(table.cell(2,2))
        #table.cell(3,2).merge(table.cell(4,2)2

        #semester filling
        col = table11.columns[2].cells
        col0 = table11.columns[0].cells
        i1 = 0
        for k0 in range(1,r1,4):
            col0[k0].text = dates[i1]
            col0[k0].paragraphs[0].runs[0].font.size = Pt(8)
            i1 = i1+1
        if(ns[-1]=='1'):
            for k2 in range(1,r1,2):
                col[k2].text = sem[-1]
                col[k2].paragraphs[0].runs[0].font.size = Pt(10)
            for m1 in range(1,r1,2):
                table11.cell(m1,2).merge(table11.cell(m1+1,2))
                table11.cell(m1,3).merge(table11.cell(m1+1,3))
                table11.cell(m1,4).merge(table11.cell(m1+1,4))
                table11.cell(m1,5).merge(table11.cell(m1+1,5))
                table11.cell(m1,6).merge(table11.cell(m1+1,6))
                table11.cell(m1,7).merge(table11.cell(m1+1,7))
                
        elif(ns[-1]=='2'):
            for k2 in range(1,r1+1):
                if(k2%2!=0):
                    col[k2].text = sem[-1]
                    col[k2].paragraphs[0].runs[0].font.size = Pt(10)
                else:
                    col[k2].text = sem1[-1]
                    col[k2].paragraphs[0].runs[0].font.size = Pt(10)

        for k in range(1,r1,4):
            table11.cell(k,0).merge(table11.cell(k+3,0))
        for k1 in range(1,r1,2):
            table11.cell(k1,1).merge(table11.cell(k1+1,1))


        table11.style = 'Table Grid'



        
        

        #go for new page
        doc.add_page_break()
        doc_para1 = doc.add_paragraph(t[-1]+' - Hall Chart and Attendance Cover')
        doc_para1.alignment=WD_ALIGN_PARAGRAPH.CENTER

        """table2 = doc.add_table(rows=3, cols=2)

        #adding heading for first table
        row = table2.rows[0].cells
        row[0].add_paragraph('Degree & Department: B.E - EEE')
        row[1].add_paragraph('Slot - I: 09.30 A.M to 11.00 A.M')
        row = table2.rows[1].cells
        row[0].add_paragraph('Date: '+date_s[-1]+' - '+date_e[-1])
        row[1].add_paragraph('Slot - II: 11.30 A.M to 01.00 P.M')
        row = table2.rows[2].cells
        row[1].add_paragraph('Slot - III: 02.30 P.M to 4.00 P.M')

        doc_para = doc.add_paragraph('          ')"""
        if(ns[-1]=='1'):
            doc_para = doc.add_paragraph(hall_no1[-1]+"  ("+str(len(room2))+") - Sem "+sem[-1]+" ("+str(len(room2))+") - "+slot[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1]=='2'):
            doc_para = doc.add_paragraph(hall_no1[-1]+"  ("+str(len(room2)+len(room4))+") - Sem "+sem[-1]+" ("+str(len(room2))+") "+slot[-1]+" - Sem "+sem1[-1]+" ("+str(len(room4))+")  "+slot1[-1])
            doc_para.alignment=WD_ALIGN_PARAGRAPH.CENTER

        table3 = doc.add_table(rows=10, cols=10)
        table3.allow_autofit = False

        table3.cell(0,0).merge(table3.cell(0,1))
        table3.cell(0,2).merge(table3.cell(0,3))
        table3.cell(0,4).merge(table3.cell(0,5))
        table3.cell(0,6).merge(table3.cell(0,7))
        table3.cell(0,8).merge(table3.cell(0,9))

        #adding heading
        row = table3.rows[0].cells
        p=row[0].add_paragraph('Column 1')
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[2].add_paragraph('Column 2')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[4].add_paragraph('Column 3')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[6].add_paragraph('Column 4')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[8].add_paragraph('Column 5')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER

        if(ns[-1] == '1'):
            row = table3.rows[1].cells
            p=row[0].add_paragraph('Sem '+sem[-1])
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=row[2].add_paragraph('Sem '+sem[-1])
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=row[4].add_paragraph('Sem '+sem[-1])
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=row[6].add_paragraph('Sem '+sem[-1])
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p4=row[8].add_paragraph('Sem '+sem[-1])
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif(ns[-1] == '2'):
            row = table3.rows[1].cells
            for i in range(10):
                if(i%2==0):
                    p=row[i].add_paragraph('Sem '+sem[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p=row[i].add_paragraph('Sem '+sem1[-1])
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

        #table 3 filling
        if(ns[-1] == '1'):
            flag1 = 0
            rn1 = list(room2.keys())
            if(len(room2)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(0,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn1[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn1)):
                        flag1 =1
                        break
        elif(ns[-1] == '2'):
            flag1 = 0
            rn1 = list(room2.keys())
            if(len(room2)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(0,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn1[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn1)):
                        flag1 =1
                        break
            flag1 = 0
            rn3 = list(room4.keys())
            if(len(room4)<=35):
                lim = 9
            else:
                lim = 10
            k1 = 0
            for i1 in range(1,10,2):
                col = table3.columns[i1].cells
                if(flag1):
                    break
                for j1 in range(2,lim):
                    col[j1].text = rn3[k1]
                    col[j1].paragraphs[0].runs[0].font.size = Pt(10)
                    k1 = k1+1
                    if(k1>=len(rn3)):
                        flag1 =1
                        break
            
        #table2.style = 'Table Grid'
        table3.style = 'Table Grid'

        #2nd page attendence cover
        def daterange(date1, date2):
            for n in range(int((date2 - date1).days)+1):
                yield date1 + timedelta(n)
        date1 = datetime.strptime(date_s[-1], '%d.%m.%y')
        date2 = datetime.strptime(date_e[-1], '%d.%m.%y')

        weekdays = [5, 6]
        dates = []
        for dt in daterange(date1, date2):
            if dt.weekday() not in weekdays:                    # to print only the weekdates
                dates.append(dt.strftime("%d.%m.%Y"))
        r1 = len(dates)*4

        
        doc_para2 = doc.add_paragraph("           ")
        #attendence cover
        
        table12 = doc.add_table(rows=r1+1, cols=8)

        #Adding height
        for row in table12.rows:
            row.height = Cm(0.1)

        #row heading
        row = table12.rows[0].cells
        p=row[0].add_paragraph('Date')
        #p.paragraphs[0].runs[0].font.size = Pt(10)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1=row[1].add_paragraph('Session')
        p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2=row[2].add_paragraph('Semester')
        p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p3=row[3].add_paragraph('Reg.No.of.Absentees')
        p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p4=row[4].add_paragraph('Absent')
        p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p5=row[5].add_paragraph('Present')
        p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p6=row[6].add_paragraph('Total')
        p6.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p7=row[7].add_paragraph('Signature')
        p7.alignment=WD_ALIGN_PARAGRAPH.CENTER
        #table.cell(1,2).merge(table.cell(2,2))
        #table.cell(3,2).merge(table.cell(4,2)2

        #semester filling
        col = table12.columns[2].cells
        col00 = table12.columns[0].cells
        i11 = 0
        for k00 in range(1,r1,4):
            col00[k00].text = dates[i11]
            col00[k00].paragraphs[0].runs[0].font.size = Pt(8)
            i11 = i11+1
        if(ns[-1]=='1'):
            for k2 in range(1,r1,2):
                col[k2].text = sem[-1]
                col[k2].paragraphs[0].runs[0].font.size = Pt(10)
            for m2 in range(1,r1,2):
                table12.cell(m2,2).merge(table12.cell(m2+1,2))
                table12.cell(m2,3).merge(table12.cell(m2+1,3))
                table12.cell(m2,4).merge(table12.cell(m2+1,4))
                table12.cell(m2,5).merge(table12.cell(m2+1,5))
                table12.cell(m2,6).merge(table12.cell(m2+1,6))
                table12.cell(m2,7).merge(table12.cell(m2+1,7))
        elif(ns[-1]=='2'):
            for k2 in range(1,r1+1):
                if(k2%2!=0):
                    col[k2].text = sem[-1]
                    col[k2].paragraphs[0].runs[0].font.size = Pt(10)
                else:
                    col[k2].text = sem1[-1]
                    col[k2].paragraphs[0].runs[0].font.size = Pt(10)

        for k in range(1,r1,4):
            table12.cell(k,0).merge(table12.cell(k+3,0))
        for k1 in range(1,r1,2):
            table12.cell(k1,1).merge(table12.cell(k1+1,1))


        table12.style = 'Table Grid'



        

        new_path = 'Attendance Cover Document\\'
        name = str(hall_no[-1])+" "+str(hall_no1[-1])+" "+str(t[-1])+" ("+str(date_s[-1])+") "+"Attendance Cover.docx"
        if(not os.path.exists(new_path)):
                os.makedirs(new_path)
        doc.save(new_path+name)


        
        show2()








    #starting of tkinter window
    root1 = Toplevel()
    #root1 = tkinter.Tk()
    root1.configure(background='lavender')
    root1.title("Hall Chart")
    root1.geometry("700x700")
    #root1.iconbitmap('tce.ico')

    # Heading
    w2 = tkinter.Label(root1,  text="Hall Chart", fg="black", bg="lavender")
    w2.config(font=("Elephant", 25))
    w2.grid(row=1, column=2, columnspan=2, padx=0)

    # labels
    NameLb = tkinter.Label(root1, text="Enter the exam title",anchor='w',width=40, fg="black", bg="lavender")
    NameLb.config(font=("Aharoni", 10))
    NameLb.grid(row=6, column=1, pady=10)

    NameLb1 = tkinter.Label(root1, text="Enter the starting date",anchor='w',width=40, fg="black", bg="lavender")
    NameLb1.config(font=("Aharoni", 10))
    NameLb1.grid(row=7, column=1, pady=10)

    NameLb2 = tkinter.Label(root1, text="Enter the ending date",anchor='w',width=40, fg="black", bg="lavender")
    NameLb2.config(font=("Aharoni", 10))
    NameLb2.grid(row=8, column=1, pady=10)

    NameLb3 = tkinter.Label(root1, text="Hall Name 1",anchor='w',width=40,fg="black", bg="lavender")
    NameLb3.config(font=("Aharoni", 10))
    NameLb3.grid(row=9, column=1, pady=10)

    NameLb31 = tkinter.Label(root1, text="Hall Name 2",anchor='w',width=40,fg="black", bg="lavender")
    NameLb31.config(font=("Aharoni", 10))
    NameLb31.grid(row=10, column=1, pady=10)

    NameLb4 = tkinter.Label(root1, text="UG/PG",anchor='w',width=40, fg="black", bg="lavender")
    NameLb4.config(font=("Aharoni", 10))
    NameLb4.grid(row=11, column=1, pady=10)

    NameLb5 = tkinter.Label(root1, text="Number of semester accomodated",anchor='w',width=40, fg="black", bg="lavender")
    NameLb5.config(font=("Aharoni", 10))
    NameLb5.grid(row=12, column=1, pady=10)

    NameLb6 = tkinter.Label(root1, text="Choose the Semester",anchor='w',width=40, fg="black", bg="lavender")
    NameLb6.config(font=("Aharoni", 10))
    NameLb6.grid(row=13, column=1, pady=10)

    NameLb7 = tkinter.Label(root1, text="Choose the Slot",anchor='w',width=40, fg="black", bg="lavender")
    NameLb7.config(font=("Aharoni", 10))
    NameLb7.grid(row=14, column=1, pady=10)

    NameLb8 = tkinter.Label(root1, text="Enter the respective Core Subject Code",anchor='w',width=40, fg="black", bg="lavender")
    NameLb8.config(font=("Aharoni", 10))
    NameLb8.grid(row=15, column=1, pady=10)

    NameLb9 = tkinter.Label(root1, text="Choose the Semester",anchor='w',width=40, fg="black", bg="lavender")
    NameLb9.config(font=("Aharoni", 10))
    NameLb9.grid(row=16, column=1, pady=10)

    NameLb10 = tkinter.Label(root1, text="Choose the Slot",anchor='w',width=40, fg="black", bg="lavender")
    NameLb10.config(font=("Aharoni", 10))
    NameLb10.grid(row=17, column=1, pady=10)

    NameLb11 = tkinter.Label(root1, text="Enter the respective Core Subject Code",anchor='w',width=40, fg="black", bg="lavender")
    NameLb11.config(font=("Aharoni", 10))
    NameLb11.grid(row=18, column=1, pady=10)

    FootLbb = tkinter.Label(root1, text="Version 1.0.4, Designed for EEE TCE", fg="black", bg="lavender")
    FootLbb.config(font=("Aharoni", 8))
    FootLbb.place(relx = 0.0,rely = 1.0,anchor ='sw')

    # Entry fields

    root1.hallEn = tkinter.Entry(root1, textvariable=hall)
    root1.hallEn.grid(row=9, column=2,padx=5)

    root1.hallEn1 = tkinter.Entry(root1, textvariable=hall1)
    root1.hallEn1.grid(row=10, column=2,padx=5)

    root1.sub1En = tkinter.Entry(root1, textvariable=sub1)
    root1.sub1En.grid(row=15, column=2,padx=5)

    root1.sub2En = tkinter.Entry(root1, textvariable=sub2)
    root1.sub2En.grid(row=18, column=2,padx=5)

    #Create a Calendar using DateEntry
    cal = DateEntry(root1, width= 16, background= "magenta3", foreground= "white",bd=2)
    cal.grid(row=7,column=2,padx=5)

    cal1 = DateEntry(root1, width= 16, background= "magenta3", foreground= "white",bd=2)
    cal1.grid(row=8,column=2,padx=5)

    #DropDown

    n = tkinter.StringVar()
    examtitle = ttk.Combobox(root1, width = 17, textvariable = n)

    # Adding combobox drop down list
    examtitle['values'] = ('CAT 1',
                                                    '   ',
                                                    'CAT 3',
                                                    'Semester')

    examtitle.grid(row=6, column=2, padx=5)
    examtitle.current(0)


    n1 = tkinter.StringVar()
    ug_pg = ttk.Combobox(root1, width = 17, textvariable = n1)

    # Adding combobox drop down list
    ug_pg['values'] = ('UG',
                                                    'PG')

    ug_pg.grid(row=11, column=2, padx=5)
    ug_pg.current(0)


    n2 = tkinter.StringVar()
    sem_acc = ttk.Combobox(root1, width = 17, textvariable = n2)

    # Adding combobox drop down list
    sem_acc['values'] = ('1',
                                                    '2')

    sem_acc.grid(row=12, column=2, padx=5)
    sem_acc.current(0)


    n3 = tkinter.StringVar()
    se1 = ttk.Combobox(root1, width = 17, textvariable = n3)

    # Adding combobox drop down list
    se1['values'] = ('I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII')
    se1.grid(row=13, column=2, padx=5)
    se1.current(0)


    n4 = tkinter.StringVar()
    s1 = ttk.Combobox(root1, width = 17, textvariable = n4)

    # Adding combobox drop down list
    s1['values'] = ('A', 'B')
    s1.grid(row=14, column=2, padx=5)
    s1.current(0)


    n5 = tkinter.StringVar()
    se2 = ttk.Combobox(root1, width = 17, textvariable = n5)

    # Adding combobox drop down list
    se2['values'] = ('I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII')
    se2.grid(row=16, column=2, padx=5)
    se2.current()


    n6 = tkinter.StringVar()
    s2 = ttk.Combobox(root1, width = 17, textvariable = n6)

    # Adding combobox drop down list
    s2['values'] = ('A', 'B')
    s2.grid(row=17, column=2, padx=5)
    s2.current()


    # buttons
    root1.ele = tkinter.Button(root1, text="   Elective count   ", command=ele_co,bg="bisque",fg="blue")
    root1.ele.place(x=20, y=610)
    root1.ele.config(font=("Aharoni", 12))
    
    root1.proc = tkinter.Button(root1, text="   Proceed   ", command=proceed,bg="bisque",fg="blue")
    root1.proc.place(x=450, y=610)
    root1.proc.config(font=("Aharoni", 12))

    root1.ext = tkinter.Button(root1, text="     Exit    ", command=root1.destroy,bg="bisque",fg="blue")
    root1.ext.place(x=600, y=610)
    root1.ext.config(font=("Aharoni", 12))

    root1.mainloop()


#staff incharge function

def staffincharge():
    def ask_timetable():
        messagebox.showinfo("Info","Load the Time Table Excel File")
    def ask_stafflist():
        messagebox.showinfo("Info","Load the Faculty Details Excel File")
    def complete():
        messagebox.showinfo("Success","Document created")

    Hall1 = []
    Hall2 = []
    Hall3 = []
    Hall4 = []
    Hallpg = [0]
    def doc():
        #Hall1 
        z = staff0.Hall1.get()
        Hall1.append(z)
        Hall11 = Hall1[0]
        #Hall2
        y = staff0.Hall2.get()
        Hall2.append(y)
        Hall22 = Hall2[0]
        #Hall3
        x = staff0.Hall3.get()
        Hall3.append(x)
        Hall33 = Hall3[0]
        #Hall4
        w = staff0.Hall4.get()
        Hall4.append(w)
        Hall44 = Hall4[0]
        #Hallpg
        v = staff0.Hallpg.get()
        Hallpg.append(v)
        Hallpgg = Hallpg[-1]
        ask_timetable()
        #getting time table
        tt = askopenfilename()
        #get staff excel sheet
        ask_stafflist()
        ss = askopenfilename()
        #exam dates
        examd1 = pd.read_excel(tt,usecols="B").values
        examd2  = examd1.flatten()
        #semester details
        semes = pd.read_excel(tt,usecols="F").values
        semes1 = semes.flatten()
        Odd_or_even = semes1[0]
        print("semes1")
        print(semes1)
        print("examd2")
        print(examd2)
        #slot details
        slot1 = pd.read_excel(tt,usecols="C").values
        slot2 = slot1.flatten()
        print(slot2)
        cat = slot2[1]
        #department staff list
        staff1 = pd.read_excel(ss,usecols="B").values
        staff2 = staff1.flatten()
        staff2 = [u for u in staff2 if pd.notna(u)]
        print("staff2")
        print(staff2)
        #extra duty staff list
        extrad1 = pd.read_excel(ss,usecols="C").values
        extrad2 = extrad1.flatten()
        extrad = [g for g in extrad2 if pd.notna(g)]
        #other department staff list
        otherdep1 = pd.read_excel(ss,usecols="D").values
        otherdep = otherdep1.flatten()
        otherdep = [y for y in otherdep if pd.notna(y)]
        print("otherdep")
        print(otherdep)
        #research scholar list
        research1 = pd.read_excel(ss,usecols="E").values
        research = research1.flatten()
        research = [z for z in research if pd.notna(z)]
        #one duty staff list get only one duty
        one_duty1 = pd.read_excel(ss,usecols="F").values
        one_duty = one_duty1.flatten()
        one_duty = [w for w in one_duty if pd.notna(w)]
        print("One duty staff")
        print(one_duty)
        print("research")
        print(research)
        #extrad = extrad2[1:18]
        staff3 = staff2[0:]
        for i in otherdep:
            staff3 = numpy.append(staff3,i)
        for j in research:
            staff3= numpy.append(staff3,j)
        staff4 = numpy.concatenate((staff3, staff3))
        for k in one_duty:
            staff4 = numpy.insert(staff4,17,k)
        collegef = staff2[0:]
        researchf = research[0:]
        staff5 = numpy.concatenate((staff4, collegef))
        collegefne1 = []
        print("extrad")
        print(extrad)
        for c in range(len(extrad)):
            if(extrad[c]==0 or extrad[c]<max(extrad)):
                collegefne1.append(collegef[c])
        collegefne = collegefne1[::-1]
        collegefne = numpy.concatenate((collegefne,staff2[::-1]))
        if(len(collegefne)<10):
             collegefne = numpy.concatenate((collegefne,staff2[::-1]))
        temp0 = collegefne[0]
        temp1 = collegefne[1]
        collegefne[0] = temp1
        collegefne[1] = temp0
        staff = numpy.concatenate((staff5,collegefne))
        staff = numpy.concatenate((staff5,researchf))
        print('collegefne')
        print(collegefne)
        print(researchf)
        print("staff")
        print(staff)
        slot_in1 = []
        examd3 = []
        #removing unwanted word and nan value
        for i in examd2:
            if(i!='DATE'):
                examd3.append(i)
        print(examd3)
        #removing null values
        examd3 = [x for x in examd3 if pd.notna(x)]
        #sorting in date format
        examd3.sort(key = lambda date: datetime.strptime(date, '%d-%m-%y'))
        print(examd3)
        examd = []
        temp = examd3[0]
        examd.append(examd3[0])
        for i in range(1,len(examd3)):
            if(examd3[i]==temp):
                continue
            else:
                examd.append(examd3[i])
                temp = examd3[i]
        print(examd)
        doc = docx.Document('Default.docx')
        l=0
        print(semes1)
        print(Hall11)
        print(staff)
        for i in range(len(examd)):
            temps1 = ""
            temps2 = ""
            temps3 = ""
            print(l)
            table = doc.add_table(rows=13, cols=5, style='Table Grid')
            row = table.rows[0].cells
            row1 = table.rows[1].cells
            row2 = table.rows[2].cells
            row3 = table.rows[3].cells
            row4 = table.rows[4].cells
            row5 = table.rows[5].cells
            row6 = table.rows[6].cells
            row7 = table.rows[7].cells
            row8 = table.rows[8].cells
            row9 = table.rows[9].cells
            row10 = table.rows[10].cells
            row11 = table.rows[11].cells
            row12 = table.rows[12].cells
            col = table.columns[0].cells
            col1 = table.columns[1].cells
            col2 = table.columns[2].cells
            col3 = table.columns[3].cells
            col4 = table.columns[4].cells
            row[0].text = "     Date of Invigilation    "
            row[1].text = "    Slot    "
            row[2].text = "    Sem    "
            row[3].text = "    Hall Name    "
            row[4].text = "    Name of Invigilator    "
            p=col[5].add_paragraph(examd[i])
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p1=col1[1].add_paragraph("09:30 AM to 11:00 AM")
            p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2=col1[5].add_paragraph("11:30 AM to 01:00 PM")
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p3=col1[9].add_paragraph("02:30 PM to 4:00 PM")
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            print("jloop")
            print("Len of staff")
            print(len(staff))
            print(extrad)
            print("Len of examd2")
            print(len(examd2))
            #check the sorted date list with original list to find correct date and slot to write semester details
            for j in range(len(examd2)):
                if(examd[i]==examd2[j]):
                    if(slot2[j]=="09:30AM to 11:00AM"):
                        print(semes1[j])
                        if('PG' in semes1[j]):
                            rowe1 = table.add_row().cells
                            rowe1[1].text = '09:30AM to 11:00AM'
                            rowe1[2].text = semes1[j]
                            rowe1[3].text = Hallpgg
                            rowe1[4].text = staff[l]
                            l=l+1
                            print(staff[l])
                            table.cell(1,0).merge(table.cell(13,0))
                        elif(temps1 == "09:30AM to 11:00AM"):
                            #creating four extra row
                            rowe2 = table.add_row().cells
                            rowe3 = table.add_row().cells
                            rowe4 = table.add_row().cells
                            rowe5 = table.add_row().cells
                            #adding text in each row
                            rowe2[2].text = semes1[j]
                            rowe2[3].text = Hall11
                            rowe2[4].text = staff[l]
                            l = l+1
                            rowe3[1].text = "09:30AM to 11:00AM"
                            rowe3[2].text = semes1[j]
                            rowe3[3].text = Hall22
                            rowe3[4].text = staff[l]
                            l = l+1
                            rowe4[2].text = semes1[j]
                            rowe4[3].text = Hall33
                            rowe4[4].text = staff[l]
                            l = l+1
                            rowe5[2].text = semes1[j]
                            rowe5[3].text = Hall44
                            rowe5[4].text = staff[l]
                            l = l+1
                            table.cell(1,0).merge(table.cell(-1,0))
                            table.cell(-4,1).merge(table.cell(-1,1))
                            table.cell(-8,1).merge(table.cell(-5,1))
                        else:
                            col2[1].text = semes1[j]
                            col2[2].text = semes1[j]
                            col2[3].text = semes1[j]
                            col2[4].text = semes1[j]
                            col4[1].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[2].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[3].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col4[4].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col3[1].text = Hall11
                            col3[2].text = Hall22
                            col3[3].text = Hall33
                            col3[4].text = Hall44
                        temps1 = "09:30AM to 11:00AM"
                    if(slot2[j]=='11:30AM to 01:00PM'):
                        print(semes1[j])
                        if('PG' in semes1[j]):
                            rowe1 = table.add_row().cells
                            rowe1[1].text = '11:30AM to 01:00PM'
                            rowe1[2].text = semes1[j]
                            rowe1[3].text = Hallpgg
                            rowe1[4].text = staff[l]
                            l=l+1
                            print(staff[l])
                            table.cell(1,0).merge(table.cell(13,0))
                        elif(temps2 == '11:30AM to 01:00PM'):
                            #creating four extra row
                            rowe2 = table.add_row().cells
                            rowe3 = table.add_row().cells
                            rowe4 = table.add_row().cells
                            rowe5 = table.add_row().cells
                            #adding text in each row
                            rowe2[2].text = semes1[j]
                            rowe2[3].text = Hall11
                            rowe2[4].text = staff[l]
                            l = l+1
                            rowe3[1].text = '11:30AM to 01:00PM'
                            rowe3[2].text = semes1[j]
                            rowe3[3].text = Hall22
                            rowe3[4].text = staff[l]
                            l = l+1
                            rowe4[2].text = semes1[j]
                            rowe4[3].text = Hall33
                            rowe4[4].text = staff[l]
                            l = l+1
                            rowe5[2].text = semes1[j]
                            rowe5[3].text = Hall44
                            rowe5[4].text = staff[l]
                            l = l+1
                            table.cell(1,0).merge(table.cell(-1,0))
                            table.cell(-4,1).merge(table.cell(-1,1))
                            table.cell(-8,1).merge(table.cell(-5,1))
                        else:
                            col2[5].text = semes1[j]
                            col2[6].text = semes1[j]
                            col2[7].text = semes1[j]
                            col2[8].text = semes1[j]
                            col4[5].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[6].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[7].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col4[8].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col3[5].text = Hall11
                            col3[6].text = Hall22
                            col3[7].text = Hall33
                            col3[8].text = Hall44
                        temps2 = "11:30AM to 01:00PM"
                    if(slot2[j]=='02:30PM to 4:00PM'):
                        print(semes1[j])
                        if('PG' in semes1[j]):
                            rowe1 = table.add_row().cells
                            rowe1[1].text = '02:30PM to 4:00PM'
                            rowe1[2].text = semes1[j]
                            rowe1[3].text = Hallpgg
                            rowe1[4].text = staff[l]
                            l=l+1
                            print(staff[l])
                            table.cell(1,0).merge(table.cell(13,0))
                        elif(temps3 == '02:30PM to 4:00PM'):
                            #creating four extra row
                            rowe2 = table.add_row().cells
                            rowe3 = table.add_row().cells
                            rowe4 = table.add_row().cells
                            rowe5 = table.add_row().cells
                            #adding text in each row
                            rowe2[2].text = semes1[j]
                            rowe2[3].text = Hall11
                            rowe2[4].text = staff[l]
                            l = l+1
                            rowe3[1].text = '02:30PM to 4:00PM'
                            rowe3[2].text = semes1[j]
                            rowe3[3].text = Hall22
                            rowe3[4].text = staff[l]
                            l = l+1
                            rowe4[2].text = semes1[j]
                            rowe4[3].text = Hall33
                            rowe4[4].text = staff[l]
                            l = l+1
                            rowe5[2].text = semes1[j]
                            rowe5[3].text = Hall44
                            rowe5[4].text = staff[l]
                            l = l+1
                            table.cell(1,0).merge(table.cell(-1,0))
                            table.cell(-4,1).merge(table.cell(-1,1))
                            table.cell(-8,1).merge(table.cell(-5,1))
                        else:
                            col2[9].text = semes1[j]
                            col2[10].text = semes1[j]
                            col2[11].text = semes1[j]
                            col2[12].text = semes1[j]
                            col4[9].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[10].text = staff[l]
                            l = l+1
                            print(staff[l])
                            col4[11].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col4[12].text = staff[l]
                            l=l+1
                            print(staff[l])
                            col3[9].text = Hall11
                            col3[10].text = Hall22
                            col3[11].text = Hall33
                            col3[12].text = Hall44
                        temps3 = "02:30PM to 4:00PM"
            #merging the cells for slots
            table.cell(1,0).merge(table.cell(12,0))
            table.cell(1,1).merge(table.cell(4,1))
            table.cell(5,1).merge(table.cell(8,1))
            table.cell(9,1).merge(table.cell(12,1))
            doc_para = doc.add_paragraph("         ")
            doc_para = doc.add_paragraph("         ")
            doc.add_page_break()
          
        # Now save the document to a location
        new_path = 'Invigilation Document\\'
        name = Odd_or_even+' '+cat+' '+examd[0]+' - '+examd[-1]+' '+' exam duty.docx'
        if(not os.path.exists(new_path)):
                os.makedirs(new_path)
        doc.save(new_path+name)
        print("staff len = "+str(len(staff)))
        print("len of college faculty = "+str(len(collegef)))
        print("l = "+str(l))
        print(len(staff)-len(collegef))
        limitd = len(staff)-len(collegef)
        print("len of staff5 = "+str(len(staff5)))
        collegef1 = collegef[::-1]
        if(l>limitd):
            print("-------------------------creating extra duty-------------------------------")
            extraduty = l-limitd
            doc1 = docx.Document('Default.docx')
            tablee1 = doc1.add_table(rows=1, cols=1, style='Table Grid')
            rowed = tablee1.rows[0].cells
            rowed[0].text = "Extra duty"
            for d in range(extraduty):
                rowed1 = tablee1.add_row().cells
                print(collegefne[d])
                rowed1[0].text = collegefne[d]
            new_path1 = 'Invigilation\\'
            name1 = Odd_or_even+' '+cat+' '+examd[0]+' - '+examd[-1]+' '+' extra duty.docx'
            if(not os.path.exists(new_path1)):
                    os.makedirs(new_path1)
            doc1.save(new_path1+name1)
        complete()


    staff0 = Toplevel()
    staff0.configure(background='lavender')
    staff0.title("Staff Invigilation")
    staff0.geometry("700x500")
    #staff0.iconbitmap('tce.ico')

    #Labels
    hlab1 = tkinter.Label(staff0, text="UG", fg="blue", bg="lavender")
    hlab1.config(font=("Aharoni", 30))
    hlab1.place(x=10,y=10)

    hlab2 = tkinter.Label(staff0, text="Enter the Hall Numbers", fg="black", bg="lavender")
    hlab2.config(font=("Aharoni", 20))
    hlab2.place(x=10,y=80)

    hlab4 = tkinter.Label(staff0, text="PG", fg="blue", bg="lavender")
    hlab4.config(font=("Aharoni", 30))
    hlab4.place(x=10,y=250)

    hlabpg = tkinter.Label(staff0, text="Enter the Hall Number", fg="black", bg="lavender")
    hlabpg.config(font=("Aharoni", 20))
    hlabpg.place(x=10,y=320)

    hlaLba = tkinter.Label(staff0, text="Version 1.0.4, Designed for EEE TCE", fg="black", bg="lavender")
    hlaLba.config(font=("Aharoni", 8))
    hlaLba.place(relx = 0.0,rely = 1.0,anchor ='sw')
    #Entries

    staff0.Hall1 = tkinter.Entry(staff0,textvariable=Hall1)
    staff0.Hall1.place(x=15,y=130,width=50,height=30)

    staff0.Hall2 = tkinter.Entry(staff0,textvariable=Hall2)
    staff0.Hall2.place(x=115,y=130,width=50,height=30)

    staff0.Hall3 = tkinter.Entry(staff0,textvariable=Hall3)
    staff0.Hall3.place(x=215,y=130,width=50,height=30)

    staff0.Hall4 = tkinter.Entry(staff0,textvariable=Hall4)
    staff0.Hall4.place(x=315,y=130,width=50,height=30)

    staff0.Hallpg = tkinter.Entry(staff0,textvariable=Hallpg)
    staff0.Hallpg.place(x=15,y=370,width=50,height=30)

    #buttons

    staff0.submit = tkinter.Button(staff0,text = "  Proceed  ",command=doc,bg="bisque",fg="blue")
    staff0.submit.place(x=400,y=450)
    staff0.submit.config(font=("Aharoni", 12))

    staff0.close = tkinter.Button(staff0, text="   Exit   ", command=staff0.destroy,bg="bisque",fg="blue")
    staff0.close.place(x=550,y=450)
    staff0.close.config(font=("Aharoni", 12))


    staff0.mainloop()

    










#main window

main1 = tkinter.Tk()
main1.configure(background='lavender')
main1.title("Ener - EEE TCE")
main1.geometry("600x420")

#images
logo = (Image.open("Photos\logo1.png"))
logo = logo.resize((100,100), Image.ANTIALIAS)
renderlogo = ImageTk.PhotoImage(logo)
logoimg = Label(main1, image=renderlogo)
logoimg.place(x=115, y=10)
#main1.iconbitmap('tce.ico')

#labels
title1 = tkinter.Label(main1,text="ENER - One Click for All",fg="black", bg="lavender")
title1.config(font=("Elephant", 15))
title1.place(x=230,y=35)

select1 = tkinter.Label(main1,text="Select your need",fg="black", bg="lavender")
select1.config(font=("Elephant", 15))
select1.place(x=210,y=150)

FootLba = tkinter.Label(main1, text="Version 1.0.4, Designed for EEE TCE", fg="black", bg="lavender")
FootLba.config(font=("Aharoni", 8))
FootLba.place(relx = 0.0,rely = 1.0,anchor ='sw')

#buttons
main1.attendence = tkinter.Button(main1, text="   Attendance Sheet   ", command=att,bg="bisque",fg="blue")
main1.attendence.place(x=100,y=200)
main1.attendence.config(font=("Aharoni", 12))

main1.hall = tkinter.Button(main1, text="          Hall Chart          ", command=hallchart1,bg="bisque",fg="blue")
main1.hall.place(x=320,y=200)
main1.hall.config(font=("Aharoni", 12))

main1.invi = tkinter.Button(main1, text="     Staff Invigilation     ", command=staffincharge,bg="bisque",fg="blue")
main1.invi.place(x=205,y=250)
main1.invi.config(font=("Aharoni", 12))

main1.moreinfo = tkinter.Button(main1, text="          More Info          ", command=info1,bg="bisque",fg="blue")
main1.moreinfo.config(font=("Aharoni", 8))
main1.moreinfo.place(relx = 0.82,rely = 0.05,anchor ='sw')
main1.mainloop()
